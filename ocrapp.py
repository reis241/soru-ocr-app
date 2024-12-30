import pytesseract
from PIL import Image
from tkinter import Tk, filedialog, messagebox, Toplevel, Label, StringVar
from tkinter.ttk import Progressbar
from docx import Document

# Tesseract'ın kurulu olduğu yolu belirtin
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"  # Tesseract kurulum dizini

# OCR işlemini gerçekleştiren fonksiyon
def perform_ocr(image_path):
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image, lang="tur")  # 'eng' veya 'tur'
        return text
    except Exception as e:
        return f"Hata: {str(e)}"

# Progress bar ile toplu OCR işlemi
def select_files_and_ocr():
    # Tkinter dosya seçme arayüzünü başlat
    Tk().withdraw()
    file_paths = filedialog.askopenfilenames(
        title="Bir veya daha fazla resim dosyası seçin",
        filetypes=[("Resim Dosyaları", "*.png *.jpg *.jpeg *.bmp")]
    )
    if file_paths:
        # İlerleme penceresini oluştur
        progress_window = Toplevel()
        progress_window.title("İlerleme Durumu")
        Label(progress_window, text="OCR işlemi devam ediyor...").pack(pady=10)
        progress_var = StringVar()
        progress_bar = Progressbar(progress_window, length=300, mode="determinate")
        progress_bar.pack(pady=10)
        progress_label = Label(progress_window, textvariable=progress_var)
        progress_label.pack(pady=5)
        progress_bar["maximum"] = len(file_paths)

        # Word belgesini oluştur
        document = Document()
        document.add_heading("OCR Sonuçları", level=1)

        for idx, file_path in enumerate(file_paths):
            try:
                # Her bir görsel için OCR işlemi yap
                ocr_result = perform_ocr(file_path)
                document.add_heading(f"Dosya: {file_path.split('/')[-1]}", level=2)
                document.add_paragraph(ocr_result)
            except Exception as e:
                document.add_paragraph(f"{file_path}: Hata oluştu ({str(e)})")

            # İlerleme çubuğunu güncelle
            progress_bar["value"] = idx + 1
            progress_var.set(f"%{int(((idx + 1) / len(file_paths)) * 100)} tamamlandı")
            progress_window.update()

        # Word belgesini kaydet
        save_path = filedialog.asksaveasfilename(
            title="OCR sonuçlarını kaydetmek için bir Word dosyası seçin",
            defaultextension=".docx",
            filetypes=[("Word Belgesi", "*.docx")]
        )
        if save_path:
            document.save(save_path)
            messagebox.showinfo("Başarılı", f"OCR sonuçları {save_path} dosyasına kaydedildi!")
        else:
            messagebox.showwarning("Uyarı", "Kaydedilecek dosya seçilmedi.")
        
        # İlerleme penceresini kapat
        progress_window.destroy()
    else:
        messagebox.showwarning("Uyarı", "Hiçbir dosya seçilmedi.")

# Uygulamayı başlat
if __name__ == "__main__":
    select_files_and_ocr()
